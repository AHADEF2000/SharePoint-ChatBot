import os
from flask import Flask, request, jsonify, render_template
from flask_cors import CORS
from dotenv import load_dotenv
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from langchain.embeddings import OpenAIEmbeddings
from langchain.chat_models import ChatOpenAI
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory

# Load environment variables
load_dotenv()

# === Step 1: Extract text from documents ===
def extract_text_from_files(folder_path="docs"):
    texts = []
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if filename.endswith(".pdf"):
            try:
                doc = fitz.open(file_path)
                text = "\n".join([page.get_text() for page in doc])
                texts.append(Document(page_content=text, metadata={"source": filename}))
            except Exception as e:
                print(f"❌ Error reading {filename}: {e}")
        elif filename.endswith(".docx"):
            try:
                doc = DocxDocument(file_path)
                text = "\n".join([p.text for p in doc.paragraphs])
                texts.append(Document(page_content=text, metadata={"source": filename}))
            except Exception as e:
                print(f"❌ Error reading {filename}: {e}")
    return texts

# === Step 2: Chunk documents ===
def chunk_documents(documents):
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    return splitter.split_documents(documents)

# === Step 3: Load or create vector DB ===
def load_or_create_vector_db(chunks, persist_dir="persist_db"):
    os.makedirs(persist_dir, exist_ok=True)
    has_index = any(
        f.endswith(".sqlite") or f.endswith(".parquet")
        for f in os.listdir(persist_dir)
    )
    if has_index:
        print("🧠 Loading existing vector DB...")
        return Chroma(persist_directory=persist_dir, embedding_function=OpenAIEmbeddings())
    else:
        print("📦 Creating new vector DB...")
        vector_db = Chroma.from_documents(chunks, OpenAIEmbeddings(), persist_directory=persist_dir)
        vector_db.persist()
        return vector_db

# === Step 4: Build QA chain ===
def build_conversational_chain(vector_db):
    memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)
    return ConversationalRetrievalChain.from_llm(
        llm=ChatOpenAI(model="gpt-3.5-turbo"),
        retriever=vector_db.as_retriever(),
        memory=memory
    )

# === Step 5: Prepare the bot ===
print("🔄 Initializing chatbot...")
docs = extract_text_from_files("docs")
chunks = chunk_documents(docs)
vector_db = load_or_create_vector_db(chunks)
qa_chain = build_conversational_chain(vector_db)
chat_history = []
print("✅ Chatbot is ready.")

# === Step 6: Set up Flask ===
app = Flask(__name__, static_folder='static', template_folder='templates')
CORS(app)

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/chat", methods=["POST"])
def chat():
    global chat_history
    user_input = request.json.get("question", "")
    print(f"🧑 User asked: {user_input}")

    if not user_input.strip():
        return jsonify({"answer": "❗يرجى كتابة سؤال."})

    try:
        result = qa_chain.invoke({
            "question": user_input,
            "chat_history": chat_history
        })
        answer = result["answer"]
        chat_history.append((user_input, answer))
        print(f"🤖 Bot replied: {answer}")
        return jsonify({"answer": answer})
    except Exception as e:
        print(f"❌ Error during QA chain: {e}")
        return jsonify({"answer": "حدث خطأ أثناء الإجابة. حاول مرة أخرى."})

# === Step 7: Run with Azure-compatible port ===
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))
    app.run(host="0.0.0.0", port=port)
